"""Build a simple Word .docx from document title and body text."""

from __future__ import annotations

import io
from typing import Optional


def build_docx_bytes(title: str, body: str) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc
    doc = Document()
    if title:
        doc.add_heading(title[:200], level=1)
    for block in (body or "").split("\n"):
        doc.add_paragraph(block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def default_filename(title: Optional[str]) -> str:
    base = (title or "法律文书").strip() or "法律文书"
    safe = "".join(ch if ch.isalnum() or ch in "-_ " else "_" for ch in base)
    return f"{safe[:80]}.docx"
